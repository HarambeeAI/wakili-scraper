"""FastAPI routes for chat, export, citations, and document retrieval."""

from __future__ import annotations

import io
import json
import logging
import os
import uuid

import httpx
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sse_starlette.sse import EventSourceResponse

from app.config import settings
from app.auth import get_current_user
from app.database import get_db, async_session
from app.agent.state import ChatRequest, ExportRequest, ExportResponse
from app.agent.legal_agent import invoke_agent, stream_agent, get_agent
from app.agent.deep_agent import stream_deep_agent
from app.agent.drafting_agent import stream_drafting_agent
from app.agent.tools.skill_loader import detect_drafting_intent, load_skill
from app.agent.tools.documents import (
    _generate_pdf,
    _generate_docx,
    _build_legal_html,
    EXPORT_DIR,
)
from app.models import CitationRecord, Message, Thread, User, UserDocument, Canvas
from app.api.trial_emails import send_trial_nurture_email

log = logging.getLogger(__name__)

router = APIRouter()

# Scraper file server URL (same as used by the indexer)
SCRAPER_BASE_URL = os.getenv(
    "SCRAPER_BASE_URL",
    "http://wakili-scraper.railway.internal:8080",
)


def _auto_title(message: str) -> str:
    """Generate a short title from the first user message."""
    cleaned = message.strip().replace("\n", " ")
    if len(cleaned) > 80:
        return cleaned[:77] + "..."
    return cleaned


# ---------- Chat (SSE Streaming) ----------


TRIAL_CHAT_LIMIT = 10
TRIAL_DRAFT_LIMIT = 5


def _check_subscription(user: User, is_drafting: bool = False) -> dict | None:
    """Check if user can proceed. Returns error dict if blocked, None if allowed."""
    status = user.subscription_status or "trialing"

    if status == "active":
        return None

    if status == "trialing":
        if is_drafting:
            if user.trial_drafts_used >= TRIAL_DRAFT_LIMIT:
                return {
                    "error": "trial_limit_reached",
                    "type": "draft",
                    "message": f"You've used all {TRIAL_DRAFT_LIMIT} free document drafts. Subscribe to continue.",
                    "trial_drafts_remaining": 0,
                    "trial_chats_remaining": max(0, TRIAL_CHAT_LIMIT - user.trial_chats_used),
                }
        else:
            if user.trial_chats_used >= TRIAL_CHAT_LIMIT:
                return {
                    "error": "trial_limit_reached",
                    "type": "chat",
                    "message": f"You've used all {TRIAL_CHAT_LIMIT} free research queries. Subscribe to continue.",
                    "trial_chats_remaining": 0,
                    "trial_drafts_remaining": max(0, TRIAL_DRAFT_LIMIT - user.trial_drafts_used),
                }
        return None

    if status == "past_due":
        # 3-day grace period — allow access
        return None

    if status == "pending":
        return {
            "error": "payment_pending",
            "message": "Please complete your payment to access the platform.",
        }

    # cancelled or unknown
    return {
        "error": "subscription_inactive",
        "message": "Your subscription is inactive. Please subscribe to continue.",
    }


@router.post("/chat")
async def chat(
    request: ChatRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Stream a chat response with citations via SSE."""

    # Check subscription / trial limits before proceeding
    # (drafting detection happens later, so we do a preliminary chat check here;
    #  drafting-specific check happens after intent detection below)
    block = _check_subscription(user, is_drafting=False)
    if block and block.get("error") != "trial_limit_reached":
        # For non-trial blocks (pending, cancelled), block immediately
        raise HTTPException(status_code=402, detail=block)

    thread_id = request.thread_id or uuid.uuid4().hex
    is_new_thread = not request.thread_id

    # Persist the user message
    try:
        tid = uuid.UUID(thread_id)
    except ValueError:
        tid = uuid.uuid4()
        thread_id = tid.hex

    if is_new_thread:
        thread = Thread(id=tid, user_id=user.id, title=_auto_title(request.message))
        db.add(thread)
        await db.flush()
    else:
        result = await db.execute(select(Thread).where(Thread.id == tid, Thread.user_id == user.id))
        thread = result.scalar_one_or_none()
        if not thread:
            raise HTTPException(status_code=404, detail="Thread not found")

    # Save user message (with attached file metadata + text for thread persistence)
    attached = None
    if request.documents:
        attached = [{"filename": doc.filename, "text": doc.text} for doc in request.documents]
    user_msg = Message(thread_id=tid, role="user", content=request.message, attached_files=attached)
    db.add(user_msg)
    await db.commit()

    async def event_generator():
        full_content = ""
        all_citations = []
        plan_steps: list[str] = []
        completed_steps = 0

        # Build history dicts for the agent
        history = [{"role": h.role, "content": h.content} for h in request.history] if request.history else []

        # Collect document context: new attachments + docs from last 10 messages
        doc_sections = []
        seen_filenames = set()

        # 1. New documents from this message
        if request.documents:
            for doc in request.documents:
                doc_sections.append(
                    f"--- ATTACHED DOCUMENT: {doc.filename} ---\n{doc.text}\n--- END OF {doc.filename} ---"
                )
                seen_filenames.add(doc.filename)

        # 2. Documents from the last 10 messages in this thread (matching history window)
        if not is_new_thread:
            recent_msgs = await db.execute(
                select(Message).where(
                    Message.thread_id == tid,
                    Message.id != user_msg.id,
                    Message.attached_files.isnot(None),
                ).order_by(Message.created_at.desc()).limit(10)
            )
            for prev_msg in recent_msgs.scalars().all():
                for file_data in (prev_msg.attached_files or []):
                    fname = file_data.get("filename", "")
                    text = file_data.get("text", "")
                    if fname and text and fname not in seen_filenames:
                        doc_sections.append(
                            f"--- ATTACHED DOCUMENT: {fname} (from earlier in conversation) ---\n{text}\n--- END OF {fname} ---"
                        )
                        seen_filenames.add(fname)

        # Build augmented message with all document context
        augmented_message = request.message
        if doc_sections:
            doc_block = "\n\n".join(doc_sections)
            augmented_message = (
                f"The user has attached the following document(s) for context:\n\n"
                f"{doc_block}\n\n"
                f"User's question/instruction:\n{request.message}"
            )

        # Check for drafting intent first
        is_drafting, doc_type_label, skill_slug = detect_drafting_intent(request.message)
        log.info(f"Drafting intent: is_drafting={is_drafting}, doc_type={doc_type_label}, skill={skill_slug}")

        # Trial enforcement: check limits and increment counters
        if (user.subscription_status or "trialing") == "trialing":
            use_type = "draft" if is_drafting else "chat"
            trial_block = _check_subscription(user, is_drafting=is_drafting)
            if trial_block:
                yield {
                    "event": "trial_limit",
                    "data": json.dumps(trial_block),
                }
                return

            # Increment the appropriate counter
            if is_drafting:
                user.trial_drafts_used = (user.trial_drafts_used or 0) + 1
            else:
                user.trial_chats_used = (user.trial_chats_used or 0) + 1

            try:
                async with async_session() as trial_db:
                    trial_user = await trial_db.get(User, user.id)
                    if trial_user:
                        if is_drafting:
                            trial_user.trial_drafts_used = user.trial_drafts_used
                        else:
                            trial_user.trial_chats_used = user.trial_chats_used
                        # Send milestone nurture email if applicable
                        try:
                            send_trial_nurture_email(trial_user)
                        except Exception as email_err:
                            log.error(f"Trial email error: {email_err}")
                        await trial_db.commit()
            except Exception as e:
                log.error(f"Failed to update trial counter: {e}")

            # Emit usage update so frontend can update the banner
            yield {
                "event": "trial_usage",
                "data": json.dumps({
                    "type": "trial_usage",
                    "trial_chats_remaining": max(0, TRIAL_CHAT_LIMIT - (user.trial_chats_used or 0)),
                    "trial_drafts_remaining": max(0, TRIAL_DRAFT_LIMIT - (user.trial_drafts_used or 0)),
                }),
            }

        # Check if this thread already has a canvas (refinement flow)
        # Use separate session to avoid polluting the main db session
        existing_canvas = None
        existing_canvas_data = None
        if not is_new_thread:
            async with async_session() as check_db:
                canvas_result = await check_db.execute(
                    select(Canvas).where(Canvas.thread_id == tid, Canvas.user_id == user.id)
                )
                c = canvas_result.scalar_one_or_none()
                if c:
                    # Detach data before session closes
                    existing_canvas_data = {
                        "skill_slug": c.skill_slug,
                        "intake_answers": c.intake_answers or {},
                        "sections": c.sections or [],
                    }

        if is_drafting or existing_canvas_data:
            # Load skill
            skill = load_skill(skill_slug) if skill_slug else None

            if existing_canvas_data:
                # Refinement: use existing canvas data
                skill = load_skill(existing_canvas_data["skill_slug"]) if existing_canvas_data["skill_slug"] else skill
                agent_stream = stream_drafting_agent(
                    message=augmented_message,
                    thread_id=thread_id,
                    skill_slug=existing_canvas_data["skill_slug"] or skill_slug,
                    skill=skill or {},
                    web_enabled=request.web_enabled,
                    intake_answers=existing_canvas_data["intake_answers"],
                    existing_sections=existing_canvas_data["sections"],
                    is_refinement=True,
                    history=history,
                )
            elif skill:
                log.info(f"Drafting agent activated: skill={skill_slug}, title={skill.get('title')}")
                # New drafting task: create canvas record using separate session
                async with async_session() as canvas_db:
                    new_canvas = Canvas(
                        thread_id=tid,
                        user_id=user.id,
                        doc_type=skill.get("doc_type", ""),
                        title=skill.get("title", doc_type_label),
                        skill_slug=skill_slug,
                        status="intake",
                    )
                    canvas_db.add(new_canvas)
                    await canvas_db.commit()

                agent_stream = stream_drafting_agent(
                    message=augmented_message,
                    thread_id=thread_id,
                    skill_slug=skill_slug,
                    skill=skill,
                    web_enabled=request.web_enabled,
                    history=history,
                )
            else:
                # Drafting detected but no skill found — fall through to standard agent
                log.warning(f"Drafting intent detected but no skill loaded: slug={skill_slug}")
                agent_stream = stream_agent(
                    message=augmented_message,
                    thread_id=thread_id,
                    web_enabled=request.web_enabled,
                    history=history,
                )
        elif request.deep_research:
            agent_stream = stream_deep_agent(
                message=augmented_message,
                thread_id=thread_id,
                web_enabled=request.web_enabled,
                history=history,
            )
        else:
            agent_stream = stream_agent(
                message=augmented_message,
                thread_id=thread_id,
                web_enabled=request.web_enabled,
                history=history,
            )

        canvas_sections = []

        async for event in agent_stream:
            evt_type = event["type"]
            if evt_type not in ("token",):
                log.info(f"SSE event: {evt_type}")
            if evt_type == "token":
                full_content += event.get("content", "")
            elif evt_type == "citations":
                all_citations.extend(event.get("citations", []))
            elif evt_type == "plan" and event.get("steps"):
                plan_steps = event["steps"]
                completed_steps = 0
                doc_note = ""
                if event.get("document_requested"):
                    fmt = (event.get("document_format") or "pdf").upper()
                    doc_note = f"\n> A **{fmt}** document will be generated.\n"
                checklist = (
                    "\n".join(f"- [ ] {s}" for s in plan_steps)
                    + doc_note
                    + "\n\n---\n\n"
                )
                full_content += checklist
            elif evt_type == "step_complete":
                if completed_steps < len(plan_steps):
                    step_text = plan_steps[completed_steps]
                    full_content = full_content.replace(
                        f"- [ ] {step_text}",
                        f"- [x] {step_text}",
                    )
                completed_steps += 1
            elif evt_type == "canvas_section":
                canvas_sections.append({
                    "id": event.get("section_id", ""),
                    "heading": event.get("heading", ""),
                    "content": event.get("content", ""),
                    "order": event.get("order", 0),
                    "last_edited_by": "agent",
                    "confidence": event.get("confidence", ""),
                    "enrichment_prompt": event.get("enrichment_prompt", ""),
                    "topic_assessments": event.get("topic_assessments", []),
                    "comment": event.get("comment", ""),
                })
            elif evt_type == "canvas_complete":
                # Persist sections to canvas record using separate session
                if event.get("sections"):
                    canvas_sections = event["sections"]
                try:
                    async with async_session() as canvas_db:
                        canvas_result = await canvas_db.execute(
                            select(Canvas).where(Canvas.thread_id == tid)
                        )
                        canvas_obj = canvas_result.scalar_one_or_none()
                        if canvas_obj:
                            canvas_obj.sections = canvas_sections
                            canvas_obj.version = event.get("version", canvas_obj.version + 1)
                            canvas_obj.status = "drafting"
                            await canvas_db.commit()
                except Exception as e:
                    log.error(f"Failed to persist canvas sections: {e}")
            elif evt_type == "canvas_update":
                # Persist individual section update using separate session
                try:
                    async with async_session() as canvas_db:
                        canvas_result = await canvas_db.execute(
                            select(Canvas).where(Canvas.thread_id == tid)
                        )
                        canvas_obj = canvas_result.scalar_one_or_none()
                        if canvas_obj and canvas_obj.sections:
                            sections = list(canvas_obj.sections)
                            for i, s in enumerate(sections):
                                if s["id"] == event.get("section_id"):
                                    updated = {
                                        **s,
                                        "content": event.get("content", s["content"]),
                                        "last_edited_by": "agent",
                                    }
                                    if event.get("heading"):
                                        updated["heading"] = event["heading"]
                                    sections[i] = updated
                                    break
                            canvas_obj.sections = sections
                            canvas_obj.version += 1
                            await canvas_db.commit()
                except Exception as e:
                    log.error(f"Failed to persist canvas update: {e}")
            elif evt_type == "canvas_intake":
                # Persist intake answers using separate session
                try:
                    async with async_session() as canvas_db:
                        canvas_result = await canvas_db.execute(
                            select(Canvas).where(Canvas.thread_id == tid)
                        )
                        canvas_obj = canvas_result.scalar_one_or_none()
                        if canvas_obj:
                            canvas_obj.intake_answers = event.get("answers", {})
                            canvas_obj.status = "researching"
                            await canvas_db.commit()
                except Exception as e:
                    log.error(f"Failed to persist intake answers: {e}")
            elif evt_type == "canvas_comment":
                # Persist co-drafter comment to section
                try:
                    async with async_session() as canvas_db:
                        canvas_result = await canvas_db.execute(
                            select(Canvas).where(Canvas.thread_id == tid)
                        )
                        canvas_obj = canvas_result.scalar_one_or_none()
                        if canvas_obj and canvas_obj.sections:
                            sections = list(canvas_obj.sections)
                            for i, s in enumerate(sections):
                                if s["id"] == event.get("section_id"):
                                    sections[i] = {**s, "comment": event.get("comment", "")}
                                    break
                            canvas_obj.sections = sections
                            await canvas_db.commit()
                except Exception as e:
                    log.error(f"Failed to persist canvas comment: {e}")
            elif evt_type == "canvas_assessment":
                # Persist confidence + enrichment_prompt to section
                try:
                    async with async_session() as canvas_db:
                        canvas_result = await canvas_db.execute(
                            select(Canvas).where(Canvas.thread_id == tid)
                        )
                        canvas_obj = canvas_result.scalar_one_or_none()
                        if canvas_obj and canvas_obj.sections:
                            sections = list(canvas_obj.sections)
                            for i, s in enumerate(sections):
                                if s["id"] == event.get("section_id"):
                                    sections[i] = {
                                        **s,
                                        "confidence": event.get("confidence", "medium"),
                                        "enrichment_prompt": event.get("enrichment_prompt", ""),
                                        "topic_assessments": event.get("topic_assessments", []),
                                    }
                                    break
                            canvas_obj.sections = sections
                            await canvas_db.commit()
                except Exception as e:
                    log.error(f"Failed to persist canvas assessment: {e}")

            yield {"event": evt_type, "data": json.dumps(event)}

        # Persist assistant message + citations using a fresh session
        # (the main db session may be in a dirty state from canvas ops / selects)
        try:
            async with async_session() as persist_db:
                assistant_msg = Message(thread_id=tid, role="assistant", content=full_content)
                persist_db.add(assistant_msg)
                await persist_db.flush()

                for cite in all_citations:
                    persist_db.add(CitationRecord(
                        message_id=assistant_msg.id,
                        citation_id=str(cite.get("id", "")),
                        data=cite,
                    ))
                await persist_db.commit()
        except Exception as e:
            log.error(f"Failed to persist assistant message: {e}")

    return EventSourceResponse(event_generator())


@router.post("/chat/sync")
async def chat_sync(
    request: ChatRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Non-streaming chat endpoint for simpler clients."""

    thread_id = request.thread_id or uuid.uuid4().hex
    is_new_thread = not request.thread_id

    try:
        tid = uuid.UUID(thread_id)
    except ValueError:
        tid = uuid.uuid4()
        thread_id = tid.hex

    if is_new_thread:
        thread = Thread(id=tid, user_id=user.id, title=_auto_title(request.message))
        db.add(thread)
        await db.flush()

    attached = [{"filename": doc.filename, "text": doc.text} for doc in request.documents] if request.documents else None
    user_msg = Message(thread_id=tid, role="user", content=request.message, attached_files=attached)
    db.add(user_msg)
    await db.flush()

    # Collect ALL document context: new attachments + previous thread attachments
    doc_sections = []
    seen_filenames = set()

    if request.documents:
        for doc in request.documents:
            doc_sections.append(
                f"--- ATTACHED DOCUMENT: {doc.filename} ---\n{doc.text}\n--- END OF {doc.filename} ---"
            )
            seen_filenames.add(doc.filename)

    if not is_new_thread:
        recent_msgs = await db.execute(
            select(Message).where(
                Message.thread_id == tid,
                Message.id != user_msg.id,
                Message.attached_files.isnot(None),
            ).order_by(Message.created_at.desc()).limit(10)
        )
        for prev_msg in recent_msgs.scalars().all():
            for file_data in (prev_msg.attached_files or []):
                fname = file_data.get("filename", "")
                text = file_data.get("text", "")
                if fname and text and fname not in seen_filenames:
                    doc_sections.append(
                        f"--- ATTACHED DOCUMENT: {fname} (from earlier in conversation) ---\n{text}\n--- END OF {fname} ---"
                    )
                    seen_filenames.add(fname)

    augmented_message = request.message
    if doc_sections:
        doc_block = "\n\n".join(doc_sections)
        augmented_message = (
            f"The user has attached the following document(s) for context:\n\n"
            f"{doc_block}\n\n"
            f"User's question/instruction:\n{request.message}"
        )

    result = invoke_agent(
        message=augmented_message,
        thread_id=thread_id,
        web_enabled=request.web_enabled,
    )

    # Persist assistant response
    assistant_msg = Message(thread_id=tid, role="assistant", content=result["content"])
    db.add(assistant_msg)
    await db.flush()

    for cite in result.get("citations", []):
        db.add(CitationRecord(
            message_id=assistant_msg.id,
            citation_id=str(cite.get("id", "")),
            data=cite,
        ))

    await db.commit()

    return {
        "content": result["content"],
        "citations": result["citations"],
        "thread_id": thread_id,
    }


# ---------- Export ----------


@router.post("/export")
async def export_conversation(
    request: ExportRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a conversation thread as PDF or DOCX."""

    # Trial users cannot export
    if user.subscription_status == "trialing":
        raise HTTPException(
            status_code=403,
            detail="Document export is a premium feature. Please upgrade to Pro or Team."
        )

    # Load from database
    try:
        tid = uuid.UUID(request.thread_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid thread ID")

    result = await db.execute(select(Thread).where(Thread.id == tid, Thread.user_id == user.id))
    thread = result.scalar_one_or_none()
    if not thread:
        raise HTTPException(status_code=404, detail="Thread not found")

    msg_result = await db.execute(
        select(Message).where(Message.thread_id == tid).order_by(Message.created_at)
    )
    messages = msg_result.scalars().all()

    if not messages:
        raise HTTPException(status_code=404, detail="No messages in thread")

    # Collect citations from all assistant messages
    all_citations = []
    for msg in messages:
        if msg.role == "assistant":
            cite_result = await db.execute(
                select(CitationRecord).where(CitationRecord.message_id == msg.id)
            )
            for cr in cite_result.scalars().all():
                all_citations.append(cr.data)

    sections = []
    for msg in messages:
        if msg.role == "user":
            sections.append({"heading": "Query", "level": 3, "content": msg.content})
        elif msg.role == "assistant" and msg.content:
            sections.append({"heading": "Analysis", "level": 3, "content": msg.content})

    file_id = uuid.uuid4().hex[:8]
    title = "Legal Research Export"
    user_export_dir = os.path.join(EXPORT_DIR, str(user.id))

    if request.format == "pdf":
        filename = f"lawlyfy_export_{file_id}.pdf"
        html = _build_legal_html(title, sections, all_citations)
        _generate_pdf(html, filename, export_dir=user_export_dir)
    elif request.format == "docx":
        filename = f"lawlyfy_export_{file_id}.docx"
        _generate_docx(title, sections, all_citations, filename, export_dir=user_export_dir)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {request.format}")

    return ExportResponse(
        download_url=f"/exports/{filename}",
        format=request.format,
        filename=filename,
    )


# ---------- Document Download ----------


@router.get("/exports/{filename}")
async def download_export(
    filename: str,
    user: User = Depends(get_current_user),
):
    """Download an exported document (user-scoped)."""
    import re

    # Validate filename to prevent path traversal
    if not re.match(r'^lawlyfy_(export|canvas)_[a-f0-9]{8}\.(pdf|docx|pptx)$', filename):
        raise HTTPException(status_code=400, detail="Invalid filename")

    # Serve from user-scoped directory
    filepath = os.path.join(EXPORT_DIR, str(user.id), filename)

    # Fallback: check legacy flat directory for older exports
    if not os.path.exists(filepath):
        legacy_path = os.path.join(EXPORT_DIR, filename)
        if os.path.exists(legacy_path):
            filepath = legacy_path
        else:
            raise HTTPException(status_code=404, detail="File not found")

    media_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    }
    ext = os.path.splitext(filename)[1]
    media_type = media_types.get(ext, "application/octet-stream")

    return FileResponse(filepath, media_type=media_type, filename=filename)


# ---------- Document Upload (text extraction + persistence) ----------

MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10 MB per file (infrastructure limit)
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/data/uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Token budget: 150k tokens (~600k chars at ~4 chars/token)
TOTAL_TOKEN_BUDGET = 150_000
CHARS_PER_TOKEN = 4  # conservative estimate for English text


def _estimate_tokens(text: str) -> int:
    return len(text) // CHARS_PER_TOKEN


def _extract_text_from_pdf(data: bytes) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _extract_text_from_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text_from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


@router.post("/upload")
async def upload_documents(
    files: list[UploadFile] = File(...),
    used_tokens: int = Query(0, description="Tokens already used by previously attached docs"),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload documents for text extraction and persistent storage.

    No fixed file count limit. Uses a total token budget of 150k tokens.
    Each file is validated against the remaining budget. Accepted: PDF, DOCX, TXT.
    Max 10 MB per file (infrastructure limit).
    """
    remaining_budget = TOTAL_TOKEN_BUDGET - used_tokens
    results = []
    tokens_consumed_this_batch = 0

    for file in files:
        # Validate extension
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            results.append({
                "filename": file.filename,
                "error": f"Unsupported file type: {ext}. Accepted: PDF, DOCX, TXT.",
                "text": "",
                "token_estimate": 0,
            })
            continue

        # Read file content with size check
        data = await file.read()
        if len(data) > MAX_UPLOAD_SIZE:
            results.append({
                "filename": file.filename,
                "error": f"File exceeds 10 MB upload limit.",
                "text": "",
                "token_estimate": 0,
            })
            continue

        # Extract text
        try:
            if ext == ".pdf":
                text = _extract_text_from_pdf(data)
            elif ext == ".docx":
                text = _extract_text_from_docx(data)
            elif ext == ".txt":
                text = _extract_text_from_txt(data)
            else:
                text = ""

            if not text.strip():
                results.append({
                    "filename": file.filename,
                    "error": "No readable text found in document.",
                    "text": "",
                    "token_estimate": 0,
                })
                continue

            # Check token budget
            token_est = _estimate_tokens(text)
            budget_left = remaining_budget - tokens_consumed_this_batch

            if token_est > budget_left:
                results.append({
                    "filename": file.filename,
                    "error": (
                        f"Document too large for remaining context budget. "
                        f"This file needs ~{token_est:,} tokens but only "
                        f"~{budget_left:,} of {TOTAL_TOKEN_BUDGET:,} tokens remain."
                    ),
                    "text": "",
                    "token_estimate": token_est,
                })
                continue

            tokens_consumed_this_batch += token_est

            # Persist file to volume
            file_id = uuid.uuid4().hex[:12]
            safe_name = f"{file_id}_{file.filename}"
            storage_path = os.path.join(UPLOAD_DIR, str(user.id), safe_name)
            os.makedirs(os.path.dirname(storage_path), exist_ok=True)
            with open(storage_path, "wb") as f:
                f.write(data)

            # Save metadata to DB
            doc_record = UserDocument(
                user_id=user.id,
                filename=file.filename,
                storage_path=storage_path,
                file_size=len(data),
                token_estimate=token_est,
            )
            db.add(doc_record)

            results.append({
                "filename": file.filename,
                "error": "",
                "text": text,
                "size": len(data),
                "token_estimate": token_est,
            })

        except Exception as e:
            log.error(f"Text extraction failed for {file.filename}: {e}")
            results.append({
                "filename": file.filename,
                "error": f"Failed to extract text: {str(e)}",
                "text": "",
                "token_estimate": 0,
            })

    # Commit all doc records
    try:
        await db.commit()
    except Exception as e:
        log.error(f"Failed to persist document records: {e}")

    return {
        "files": results,
        "total_token_budget": TOTAL_TOKEN_BUDGET,
        "tokens_used": used_tokens + tokens_consumed_this_batch,
        "tokens_remaining": TOTAL_TOKEN_BUDGET - used_tokens - tokens_consumed_this_batch,
    }


# ---------- Document Retrieval (PDF proxy from scraper) ----------


@router.get("/documents/pdf")
async def get_document_pdf(
    local_path: str = Query(..., description="PDF path on scraper volume"),
    user: User = Depends(get_current_user),
):
    """Proxy a PDF from the scraper's file server for in-browser viewing."""
    import re

    if not local_path or ".." in local_path or not re.match(r'^[a-zA-Z0-9][a-zA-Z0-9_\-./]*\.pdf$', local_path):
        raise HTTPException(status_code=400, detail="Invalid document path")

    url = f"{SCRAPER_BASE_URL}/{local_path}"
    log.info(f"Proxying PDF from {url}")

    async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
        try:
            resp = await client.get(url)
            resp.raise_for_status()
        except httpx.HTTPStatusError as e:
            raise HTTPException(
                status_code=e.response.status_code,
                detail=f"PDF not found on scraper: {local_path}",
            )
        except httpx.HTTPError as e:
            raise HTTPException(status_code=502, detail=f"Cannot reach scraper: {e}")

    filename = local_path.split("/")[-1] if "/" in local_path else local_path

    return StreamingResponse(
        io.BytesIO(resp.content),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'inline; filename="{filename}"',
            "Content-Length": str(len(resp.content)),
            "Cache-Control": "public, max-age=86400",
        },
    )


@router.get("/documents/pages")
async def get_document_pages(
    local_path: str = Query(..., description="PDF path on scraper volume"),
    user: User = Depends(get_current_user),
):
    """Extract per-page text from a PDF for the document viewer."""
    import re

    if not local_path or ".." in local_path or not re.match(r'^[a-zA-Z0-9][a-zA-Z0-9_\-./]*\.pdf$', local_path):
        raise HTTPException(status_code=400, detail="Invalid document path")

    url = f"{SCRAPER_BASE_URL}/{local_path}"

    async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
        try:
            resp = await client.get(url)
            resp.raise_for_status()
        except httpx.HTTPStatusError:
            raise HTTPException(status_code=404, detail="PDF not found")
        except httpx.HTTPError as e:
            raise HTTPException(status_code=502, detail=f"Cannot reach scraper: {e}")

    try:
        import pdfplumber

        pages_data = []
        with pdfplumber.open(io.BytesIO(resp.content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages_data.append({
                    "page_number": i + 1,
                    "text": text,
                })

        return {
            "total_pages": len(pages_data),
            "pages": pages_data,
            "local_path": local_path,
        }
    except Exception as e:
        log.error(f"PDF page extraction failed for {local_path}: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract PDF pages")
